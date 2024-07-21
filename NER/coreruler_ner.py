from transformers import AutoTokenizer, AutoModelForTokenClassification
from transformers import pipeline
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document
import re
import json
from coreruler import ClientFactory
from coreruler.data_abstractions.query import Query, MapRefineQuery
from tqdm import tqdm

MAX_NER_LENGTH = 256
tokenizer = AutoTokenizer.from_pretrained("/domino/datasets/local/NER_testing/bert-large-NER")

SPLITTER_KWARGS = {
    "chunk_size": MAX_NER_LENGTH,
    "chunk_overlap": 0,
    "separators": ["\n\n", "\n", ".", ";", ",", " ", ""],
    "is_separator_regex": False,
}

def process_articles(input_path):
    # Load document
    doc_article = Document(input_path)
    splitter = RecursiveCharacterTextSplitter.from_huggingface_tokenizer(tokenizer, **SPLITTER_KWARGS)
    return [_ for split in map(lambda x: splitter.split_text(x.text), doc_article.paragraphs) for _ in split]

def preprocess_text(text):
    text = text.replace("\xa0", " ")
    if re.match(r"^\([A-Z]\)|^", text):
        text = " ".join(text.split()[1:])
    return text

model = AutoModelForTokenClassification.from_pretrained("/domino/datasets/local/NER_testing/bert-large-NER").cuda()

nlp = pipeline("ner", model=model, tokenizer=tokenizer)

def aggregate_ner_results(sentence_entities):
    if sentence_entities:
        result = [sentence_entities[0]]
        for entity in sentence_entities[1:]:
            word, end, score = entity["word"], entity["end"], entity["score"]
            last_pick = result[-1]
            if (entity["start"] - last_pick["end"]) < 1:
                last_pick["word"] += (word[2:] if word.startswith("##") else " " + word)
                last_pick["end"] = end
                last_pick["score"] *= score
            else:
                result.append(entity.copy())
        return result
    return []


map_prompt = """Context information is below.
--------------------
{article}
--------------------
Given the context information and no prior knowledge, answer concisely the question: {question}"""

refine_prompt = """The original question is as follows: {question}
We have provided an existing answer: {response}
We have the opportunity to refine the existing answer
(only if needed) with some more context below.
------------
{article}
------------
Given the new context, refine the original answer to better answer the question.
You must provide a very concise and short response, either original answer or refined answer."""

llm = ClientFactory.createModel(model="mistral-66k")
embedder = ClientFactory.createModel(model="ada")
assessor = Query(
    """Given this response to a question, can you say if this response is closer to a 'yes' or a 'no' ?
Question: {question}\nResponse: {response}.""",
    llm,
)

def output_parser(text):
    iterator = map(
        lambda x: re.sub(r"\W", "", x), text.replace("</s>", "").lower().split()
    )
    try:
        return next(filter(lambda _: _ in ["no", "yes", "missing"], iterator))
    except:
        print(assertion)

def map_refine_summary(entity, chunks):
    question = f"Provide a very concise and short description about {entity} based on the article."
    query = MapRefineQuery(map_prompt, refine_prompt, llm, refined_input_name="article")
    return query.invoke({"question": question, "article": chunks})


CHUNK_LENGTH = 14000

SPLITTER_SUMMARIZATION_KWARGS = {
    "chunk_size": CHUNK_LENGTH,
    "chunk_overlap": 0,
    "separators": ["\n\n", "\n", ".", ";", ",", " ", ""],
    "is_separator_regex": False,
    "length_function": len
}

def process_articles_summarization(input_path):
    # Load document
    doc_article = Document(input_path)
    splitter = RecursiveCharacterTextSplitter(**SPLITTER_SUMMARIZATION_KWARGS)
    chunks = splitter.create_documents(texts=list(map(lambda x: x.text, doc_article.paragraphs)))
    chunks = splitter.split_documents(chunks)
    chunks = [_.page_content for _ in chunks]
    new_chunks = [chunks[0]]
    for _ in chunks[1:]:
        if (len(new_chunks[-1]) + len(_)) < CHUNK_LENGTH:
            new_chunks[-1]+=_
        else:
            new_chunks.append(_)
    return new_chunks

import os

name_articles = os.listdir("inputs/")
name_articles = [s.replace('.docx', '') for s in name_articles]
print(name_articles)



# Loop through articles to get summaries and save them
# Comment this after all articles are run
name_articles = ["20231115_ICT1_A parish regime used a Cyprus middleman in bid to evade oil-industry sanctions"]
for article in name_articles:
    input_path = "inputs/" + article + ".docx"
    output_path = "outputs/" + article + "/"
    if not os.path.exists(output_path):
        os.mkdir(output_path)
    
    new_document = list(map(preprocess_text, process_articles(input_path)))
    ner_results = nlp(new_document)
    aggregated_result = [aggregate_ner_results(ner_sentence) for ner_sentence in ner_results]
    in_out_ner = tuple(zip(new_document, aggregated_result))
    aggregated_result = [_ for _ in aggregated_result if _]
    cleaned_aggregated_result = [entity for sentence_entities in aggregated_result for entity in sentence_entities if entity and any(_ in entity["entity"] for _ in ["ORG", "LOC", "PER"])]
    cleaned_aggregated_result = set(filter(lambda x: re.findall(r"\w", x) and not(x.startswith("##")), map(lambda x: x["word"], cleaned_aggregated_result)))

    documents_4_summarization = list(map(preprocess_text, process_articles_summarization(input_path)))
    
    dict_entity_summaries = {}
    for entity in tqdm(cleaned_aggregated_result):
        dict_entity_summaries[entity] = map_refine_summary(entity, documents_4_summarization)
    
    with open(output_path + "dict_entity_summaries.json", "w") as outfile:
        json.dump(dict_entity_summaries, outfile)

del nlp

# Need to install fuzzy wuzzy and python-Levenshtein
# Use networkx to build a graph
from fuzzywuzzy import fuzz, process
import networkx as nx
import numpy as np
import re
from llama_cpp import Llama, LlamaGrammar
from langchain.schema import Document
from utils import llama_cpp_check_entity_pairs